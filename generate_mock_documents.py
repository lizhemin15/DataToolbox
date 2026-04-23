#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成三份模拟数据的Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random

def set_run_font(run, font_name, size, bold=False, color=None):
    """设置文本格式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_red_line(paragraph):
    """添加红色下划线"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'FF0000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_document(date_str, news_data, transport_data):
    """创建单个文档"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 开头标注
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("★内部资料 模拟数据")
    set_run_font(run, '黑体', 14, color=(255, 0, 0))

    # 空行
    doc.add_paragraph()

    # 红头标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("国际新闻与运输情况通报")
    set_run_font(run, '方正小标宋简体', 22, bold=True, color=(255, 0, 0))

    # 发文字号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"国运字〔1999〕第{date_str[-2:]}号")
    set_run_font(run, '仿宋', 16)

    # 红线
    p = doc.add_paragraph()
    add_red_line(p)

    # 空行
    doc.add_paragraph()

    # 第一部分标题
    p = doc.add_paragraph()
    run = p.add_run("一、国际新闻动态")
    set_run_font(run, '黑体', 16, bold=True)

    # 新闻内容
    regions = ['亚洲', '欧洲', '中东北非', '美洲', '大洋洲']
    for region in regions:
        # 区域二级标题
        p = doc.add_paragraph()
        run = p.add_run(f"（一）{region}地区")
        set_run_font(run, '楷体', 15, bold=True)

        # 新闻条目
        for news in news_data[region]:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(32)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p.add_run(f"• {news}")
            set_run_font(run, '仿宋', 15)

    # 空行
    doc.add_paragraph()

    # 第二部分标题
    p = doc.add_paragraph()
    run = p.add_run("二、各区域运输保障情况")
    set_run_font(run, '黑体', 16, bold=True)

    # 运输情况内容
    for region in regions:
        # 区域二级标题
        p = doc.add_paragraph()
        run = p.add_run(f"（一）{region}地区")
        set_run_font(run, '楷体', 15, bold=True)

        # 海上运输
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        run = p.add_run(f"1. 海上运输：{transport_data[region]['海上']}")
        set_run_font(run, '仿宋', 15)

        # 空中运输
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        run = p.add_run(f"2. 空中运输：{transport_data[region]['空中']}")
        set_run_font(run, '仿宋', 15)

        # 陆上运输
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32)
        run = p.add_run(f"3. 陆上运输：{transport_data[region]['陆上']}")
        set_run_font(run, '仿宋', 15)

    # 空行
    doc.add_paragraph()
    doc.add_paragraph()

    # 落款
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("国际运输协调办公室")
    set_run_font(run, '仿宋', 15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"1999年1月{date_str[-2:]}日")
    set_run_font(run, '仿宋', 15)

    # 空行
    doc.add_paragraph()

    # 结尾标注
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("★内部资料 模拟数据")
    set_run_font(run, '黑体', 14, color=(255, 0, 0))

    return doc

# 三天的新闻数据
news_data_day1 = {
    '亚洲': [
        '日本东京地区发生6.2级地震，暂无重大人员伤亡报告，我驻日机构已启动应急预案。',
        '韩国政府宣布将加大对外贸易投资力度，预计将对我方运输业务产生积极影响。',
        '东南亚地区持续降雨，部分港口作业受到影响，建议相关运输任务做好预案。'
    ],
    '欧洲': [
        '德国法兰克福机场因天气原因临时关闭，预计影响我方3个航班的起降计划。',
        '英国议会通过新的运输安全法案，对我方在欧洲地区的运输保障提出更高要求。'
    ],
    '中东北非': [
        '沙特阿拉伯宣布将扩建吉达港，预计将提升该地区海上运输能力30%以上。',
        '埃及苏伊士运河管理局发布通告，将提高通行费用，需调整相关运输成本预算。',
        '阿联酋迪拜国际机场新增多条国际航线，为我方中东地区运输提供更多选择。'
    ],
    '美洲': [
        '美国西海岸港口工人罢工持续，洛杉矶港和长滩港作业效率下降约40%。',
        '巴西政府宣布将加强与亚太地区的贸易合作，预计将增加我方南美运输任务量。'
    ],
    '大洋洲': [
        '澳大利亚悉尼机场完成跑道扩建工程，可起降更大机型，提升运输保障能力。',
        '新西兰政府批准新的航空运输协定，为我方在大洋洲地区的空中运输提供便利。'
    ]
}

news_data_day2 = {
    '亚洲': [
        '印度尼西亚雅加达机场因火山灰影响临时关闭，我方已启动备用航线方案。',
        '越南胡志明市港口完成升级改造，吞吐能力提升25%，有利于我方东南亚运输布局。',
        '泰国政府宣布将简化货物通关手续，预计可缩短运输时间1-2天。'
    ],
    '欧洲': [
        '法国巴黎戴高乐机场新增安检通道，通行效率提升20%，有利于我方欧洲运输任务。',
        '俄罗斯远东地区遭遇暴风雪，部分陆路运输线路受阻，已启动应急预案。'
    ],
    '中东北非': [
        '以色列特拉维夫港完成扩建，新增2个集装箱码头，提升地中海东岸运输能力。',
        '土耳其伊斯坦布尔新机场正式投入运营，将成为欧亚运输的重要枢纽。'
    ],
    '美洲': [
        '加拿大温哥华港工人结束罢工，港口作业恢复正常，积压货物正在有序疏运。',
        '墨西哥政府宣布将投资50亿美元升级港口设施，为我方美洲运输提供新机遇。',
        '阿根廷布宜诺斯艾利斯机场新增直飞亚洲航线，缩短运输时间约8小时。'
    ],
    '大洋洲': [
        '澳大利亚墨尔本港完成航道疏浚工程，可停靠10万吨级货轮，提升运输能力。',
        '巴布亚新几内亚政府与我方签署运输合作协议，将新增2条定期航线。'
    ]
}

news_data_day3 = {
    '亚洲': [
        '新加坡樟宜机场连续第三年被评为全球最佳机场，为我方东南亚运输提供优质保障。',
        '马来西亚吉隆坡港口新增自动化码头，作业效率提升35%，降低运输成本。'
    ],
    '欧洲': [
        '荷兰鹿特丹港完成智能化改造，成为欧洲首个全自动化港口，大幅提升作业效率。',
        '意大利政府宣布将投资30亿欧元升级全国港口设施，预计工期3年。',
        '西班牙巴塞罗那机场新增直飞亚洲航线，为我方南欧运输提供便利。'
    ],
    '中东北非': [
        '卡塔尔多哈机场完成扩建工程，年旅客吞吐能力提升至5000万人次。',
        '摩洛哥丹吉尔港新增地中海航线，为我方北非地区运输提供新选择。'
    ],
    '美洲': [
        '美国东海岸遭遇暴风雪天气，纽约、波士顿等港口作业受到影响，预计影响持续2-3天。',
        '智利圣地亚哥机场完成跑道改造，可起降大型运输机，提升南美运输保障能力。'
    ],
    '大洋洲': [
        '斐济政府宣布将升级楠迪机场，预计将提升南太平洋地区运输能力。',
        '澳大利亚布里斯班港新增直达亚洲航线，运输时间缩短2天。'
    ]
}

# 三天的运输数据
def generate_transport_data(day):
    """生成运输数据"""
    transport_data = {}

    ships = ['远洋一号', '远洋二号', '远洋三号', '运输舰A型', '运输舰B型', '滚装船C型']
    aircrafts = ['运-8', '运-9', '伊尔-76', '安-124', '波音747货机', '空客A330货机']
    vehicles = ['东风重卡', '陕汽重卡', '北方奔驰', '斯太尔', '沃尔沃FH12', '曼恩TGX']

    regions = ['亚洲', '欧洲', '中东北非', '美洲', '大洋洲']

    for region in regions:
        transport_data[region] = {}

        # 海上运输
        ship = random.choice(ships)
        batches = random.randint(2, 5) + day
        sorties = batches * random.randint(3, 6)
        transport_data[region]['海上'] = f"出动{ship}等型号运输船{batches}批次，共计{sorties}艘次。"

        # 空中运输
        aircraft = random.choice(aircrafts)
        batches = random.randint(3, 8) + day * 2
        sorties = batches * random.randint(2, 4)
        transport_data[region]['空中'] = f"出动{aircraft}等型号运输机{batches}批次，共计{sorties}架次。"

        # 陆上运输
        vehicle = random.choice(vehicles)
        batches = random.randint(5, 12) + day * 3
        sorties = batches * random.randint(8, 15)
        transport_data[region]['陆上'] = f"出动{vehicle}等型号运输车{batches}批次，共计{sorties}车次。"

    return transport_data

# 生成三份文档
dates = ['19990101', '19990102', '19990103']
news_data_list = [news_data_day1, news_data_day2, news_data_day3]

for i, (date, news_data) in enumerate(zip(dates, news_data_list)):
    transport_data = generate_transport_data(i + 1)
    doc = create_document(date, news_data, transport_data)

    filename = f"{date}_国际新闻与运输情况通报_模拟数据.docx"
    doc.save(filename)
    print(f"已生成: {filename}")

print("\n所有文档生成完成！")
